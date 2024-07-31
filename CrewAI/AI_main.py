from crewai import Crew, Task
from textwrap import dedent
from crewai.process import Process
from AI_agent import AnalysisAgents
from AI_Tasks import AnalysisTasks
from datetime import date
import time

from dotenv import load_dotenv
load_dotenv()
from langchain.callbacks import get_openai_callback
from google.generativeai.types.generation_types import StopCandidateException
from langchain_anthropic import ChatAnthropic
from langchain_community.llms import Ollama

class PredictionCrew:
  def __init__(self):
    agents = AnalysisAgents()
    tasks = AnalysisTasks()

  def run(self):

    detection_agent = self.agents.weak_signal_detection_agent()
    analysis_agent = self.agents.signal_analysis_and_amplification_agent()
    report_agent = self.agents.report_generation_agent()

    detection_task = self.tasks.weak_signal_detection_task(detection_agent)
    analysis_task = self.tasks.signal_analysis_and_amplification_task(analysis_agent,detection_task.expected_output)
    report_task = self.tasks.report_generation_task(report_agent,analysis_task.expected_output)

    introduction_task = self.tasks.introduction(report_agent)
    conclusion_task = self.tasks.conclusion(report_agent, report_task.expected_output)

    # current_task_index = 0  # Start with the first task
    tasks_list=[introduction_task,
                detection_task,
                analysis_task,
                report_task,
                conclusion_task]
   
    crew = Crew(
      agents=[
        detection_agent,
        analysis_agent,
        report_agent
      ],
      tasks = tasks_list,
      verbose=True,
      #step_callback=test
    )

    result = crew.kickoff()
    # Cybersecurity_Analyst
    # task= self.impact_human.execute()
    print("Crew usage", crew.usage_metrics)
    print(result)

    return result


if __name__ == "__main__":

  import time
  st = time.time()

  # print("## Welcome to Analysis Crew")
  
  #financial_crew = FinancialCrew()
  #result = financial_crew.run()
  
  # with get_openai_callback() as cb:
  #   print(cb)

  try:
    # Your code that may raise the exception
    prediction_crew = PredictionCrew()
    result = prediction_crew.run()
    
  except StopCandidateException as e:
      # Handle the specific exception
      print("An exception occurred with finish reason:", e.finish_reason)
      result = prediction_crew.run()

  except TypeError as e:
    # Handle the TypeError exception here
    print("TypeError occurred:", e)
    result = prediction_crew.run()

  except IndexError as e:
    # Handle the IndexError exception here
    print("IndexError occurred:", e)
    result = prediction_crew.run()
    
  except Exception as e:
      # Handle any other exceptionsc
      print("An error occurred:", str(e))
      result = prediction_crew.run()

  
  elapsed_time = time.time() - st
  elapsed_time_minutes = time.strftime("%H:%M:%S", time.gmtime(elapsed_time))
  print('Execution time:', elapsed_time_minutes)

  report = str(result)

  from docx import Document
  from docx.shared import Pt
  from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
  from docx.oxml.ns import nsdecls
  from docx.oxml import parse_xml

  def add_custom_heading(document, text, level):
      heading = document.add_heading(level=level)
      heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
      run = heading.add_run(text)
      run.bold = True

  def add_custom_paragraph(document, text, bold=False, alignment=None):
      paragraph = document.add_paragraph()
      if alignment:
          paragraph.alignment = alignment
      run = paragraph.add_run(text)
      if bold:
          run.bold = True

  def process_paragraphs(paragraphs):
      processed_paragraphs = []
      #to process the output from gemini pro llm
      #the different characters at the start of sentences of the output from Gemini Pro
      #allows us to differentiate the sentences from headers
      for line in paragraphs.split('\n'):
          if (line.strip().startswith('**') and line.strip().endswith(':**')):
              processed_paragraphs.append(('header', 3, line.strip().replace("*","")))
          elif line.strip().startswith('##'):
              processed_paragraphs.append(('header', 2, line.strip().replace("#","")))
          elif (line.strip().startswith('**') or line.strip().endswith('**')):
              processed_paragraphs.append(('header', 4, line.strip().replace("*","")))
          elif line.strip().startswith('-') or line.strip().startswith('*') :
              processed_paragraphs.append(('normal', 0, line.strip().replace("*","")))
          elif not(line.strip()):
              pass
          else:
              #other normal lines without any special characters at the start
              #to account for outputs that are from other LLMs and not from Gemini Pro
              processed_paragraphs.append(('normal', 0, line.strip()))
      return processed_paragraphs

  doc = Document()

  # Apply default style to the document
  default_style = doc.styles['Normal']
  default_font = default_style.font
  default_font.bold = False  # Change to True if you want default text to be bold
  default_font.size = Pt(12)  # Change the font size as needed

  processed_paragraphs = process_paragraphs(report)

  for paragraph_type, level, text in processed_paragraphs:
      if paragraph_type == 'header':
          add_custom_heading(doc, text, level)
      elif paragraph_type == 'normal':
          add_custom_paragraph(doc, text, bold=False)
      # else:
      #     if '**' in text:
      #         text = text.replace('**', '')
      #         add_custom_paragraph(doc, text, bold=True)
      else:
          pass

  import datetime

  current_time = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

  doc.save("AI_Report_" + current_time + ".docx")

    


    
