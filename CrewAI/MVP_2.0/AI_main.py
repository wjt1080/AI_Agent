from crewai import Crew, Task
from textwrap import dedent
from crewai.process import Process
from AI_agent import AnalysisAgents
from AI_Tasks import AnalysisTasks
from datetime import date
import time

from dotenv import load_dotenv
load_dotenv()
from langchain.callbacks import get_openai_callback
from google.generativeai.types.generation_types import StopCandidateException
from langchain_anthropic import ChatAnthropic
from langchain_community.llms import Ollama

class FinancialCrew:
  def __init__(self):
    
    # self.introduction = "empty"
    # self.conclusion = "empty"
    
    self.impact_human = "empty"
    self.governance = "empty"
    self.mandiant = "empty"


  def run(self):
    agents = AnalysisAgents()
    tasks = AnalysisTasks()
    
    

    AI_Industry_Analyst = agents.AI_Industry_Analyst()
    Cybersecurity_Analyst = agents.Cybersecurity_Analyst()
    editor_advisor = agents.editor_advisor()
    
    
    self.impact_human = tasks.impact_human(AI_Industry_Analyst)
    self.governance = tasks.governance(Cybersecurity_Analyst)
    self.mandiant = tasks.mandiant(Cybersecurity_Analyst)

    list = [self.impact_human, self.governance, self.mandiant]
    self.introduction = tasks.introduction(editor_advisor, list)
    self.conclusion = tasks.conclusion(editor_advisor, list)


    

    # def test(x):
    #   nonlocal current_task_index
    #   nonlocal tasks_list
    #   nonlocal count 
    #   count = count + 1

    #   nonlocal flag
    #   if not flag:
    #     flag = True
    #     print("#########CALL BACK IN MAIN#######")
    #     print(x)
    #     print("#########CALL BACK END IN MAIN#######")

    #     #want the execution to proceed to the elif condition even if any LHS of the condition raises an exception
    #     #separate the conditions into individual try-except blocks.
    #     #condition_met flag becomes true if any "np tool used" cases are present
    #     condition_met = False

    #     try:
    #         if list(x.__dict__.keys())[0] == "return_values":
    #             print("Case 1")
    #             condition_met = True
    #     except (AttributeError, IndexError, KeyError):
    #         condition_met = False

    #     try:
    #         if not condition_met and x[0][0].tool == "_Exception":
    #             print("Case 2")
    #             condition_met = True
    #     except (AttributeError, IndexError, KeyError):
    #         condition_met = False

    #     try:
    #         if not condition_met and type(x[0][1][0].page_content) != str:
    #             print("Case 3")
    #             condition_met = True
    #     except (AttributeError, IndexError, KeyError):
    #         condition_met = False


    #     # Execute the loop while the flag is True
    #     while condition_met:
    #         print("START Re-execute. Search the Internet tool NOT used. ")

    #         #rerun task
    #         #need to modify this to capture each task!!!
    #         tasks_list[current_task_index].execute()

    #         print("END Re-execute")
            
    #         # Update the flag based on the conditions of the re-run
    #         try:
    #             if list(x.__dict__.keys())[0] == "return_values":
    #                 print("Case 1")
    #                 condition_met = True
    #         except (AttributeError, IndexError, KeyError):
    #             condition_met = False

    #         try:
    #             if not condition_met and x[0][0].tool == "_Exception":
    #                 print("Case 2")
    #                 condition_met = True
    #         except (AttributeError, IndexError, KeyError):
    #             condition_met = False

    #         try:
    #             if not condition_met and type(x[0][1][0].page_content) != str:
    #                 print("Case 3")
    #                 condition_met = True
    #         except (AttributeError, IndexError, KeyError):
    #             condition_met = False
    #   else:
    #     pass

    #   #count == 2 means we have finished the 2 callbacks for the current tasks
    #   #then reset flag to false for the subsequent tasks
    #   if count == 2:
    #     flag = False
    #     count = 0
    #     current_task_index+=1




    # current_task_index = 0  # Start with the first task
    tasks_list=[self.impact_human, self.governance, self.mandiant, self.introduction, self.conclusion]
    # tool_not_used = True #condition for a re-execution due to tool not used
    # no_final_answer = True
    # rerun = True
    
    # def test(x):
    #     nonlocal current_task_index
    #     nonlocal tasks_list
        
    #     #each call_back check for whether tool is used. If tool is used, likely to have the final answer (unless the case of tool used but final answer not produced happens)
    #     #tool use has a call back, final answer has a call back, so program might mistaken the second callback for the final answer as error even though tool is used beforehand in the first callback
    #     #so it has to be two flag (tool flag and final answer flag) consecutively happen to be present/TRUE to mean the final answer is the correct one. 
    #     #because there are cases where final answer is produced without tool used
    #     print("#########CALL BACK IN MAIN#######")
    #     print(x)
    #     print("#########CALL BACK END IN MAIN#######")
      
        
    #     nonlocal tool_not_used 
    #     #condition to check x
    #     if tool_not_used == True:
    #       try:
    #           if list(x.__dict__.keys())[0] == "return_values":
    #               print("Case 1")
    #               tool_not_used = True
    #       except (AttributeError, IndexError, KeyError):
    #           pass

    #       try:
    #           if x[0][0].tool == "_Exception":
    #               print("Case 2")
    #               tool_not_used = True
    #       except (AttributeError, IndexError, KeyError):
    #           pass

    #       try:
    #           if type(x[0][1][0].page_content) != str:
    #               print("Case 3")
    #               tool_not_used = True
    #       except (AttributeError, IndexError, KeyError):
    #           pass
    #     else:
    #       print("tool used")
    #       tool_not_used = False
        

    #     nonlocal no_final_answer
    #     #if condition to check x
    #     #when tool is used in previous call back then the current callback we check if there is the final answer callback present
    #     if tool_not_used == False:
    #       try:
    #         if list(x.__dict__.keys())[0] == "return_values":
    #           print("Final answer produced")
    #           no_final_answer = False
    #         else:
    #           #if the final answer is not present in this current round, reset tool_not_used to True
    #           #to trigger rerun. Because it means that something went wrong after the tool is used and the final answer is not given back
    #           tool_not_used == True
    #           print("Not the final answer without tools")
    #       except (AttributeError, IndexError, KeyError):
    #           #for unexpected cases of callback (because it is unexpected!!!), set as true to trigger rerun flag
    #           tool_not_used == True
    #           print("unexpected cases, set tool not used as True")
        
        

    #     nonlocal rerun
    #     if (tool_not_used == False and no_final_answer == False):
    #       #the best case and we can go on to the next task
    #       rerun = False
    #       current_task_index+=1
    #       print("Going to next task!")
    #     elif (tool_not_used == False and no_final_answer == True):
    #       #the case where we dont need rerun because tool is used and we just hv to check for the next callback for final answer 
    #       rerun = False
    #       print("tool used, check for next callback if it is the final answer callback!")
    #     else:
    #       #tool not used so need rerun
    #       rerun = True
    #       print("Rerun is true")

    #     if rerun:
    #         #during rerun, reset everything back to original flag so that the checking restarts
    #         print("START Re-execute. Search the Internet tool NOT used. ")
    #         tool_not_used = True
    #         no_final_answer = True
    #         tasks_list[current_task_index].execute() #will enter a new callback and the number and content of callback is not deterministic/fixed because there could be other errors besides tool not used, that occur which will trigger a rerun on its own and hence in turn trigger a new callback
            
                          
     
        
      
   
    crew = Crew(
      agents=[
        Cybersecurity_Analyst,
        AI_Industry_Analyst,
        editor_advisor
      ],
      
      tasks = tasks_list,
      verbose=True,
      #step_callback=test
    )

    result = crew.kickoff()
    # Cybersecurity_Analyst
    # task= self.impact_human.execute()
    print("Crew usage", crew.usage_metrics)
    self.ref_list = agents.ref
    return result
   



if __name__ == "__main__":

  import time
  st = time.time()

  # print("## Welcome to Analysis Crew")
  
  #financial_crew = FinancialCrew()
  #result = financial_crew.run()
  
  # with get_openai_callback() as cb:
  #   print(cb)

  try:
    # Your code that may raise the exception
    financial_crew = FinancialCrew()
    result = financial_crew.run()
    
  except StopCandidateException as e:
      # Handle the specific exception
      print("An exception occurred with finish reason:", e.finish_reason)
      result = financial_crew.run()

  except TypeError as e:
    # Handle the TypeError exception here
    print("TypeError occurred:", e)
    result = financial_crew.run()

  except IndexError as e:
    # Handle the IndexError exception here
    print("IndexError occurred:", e)
    result = financial_crew.run()
    
  except Exception as e:
      # Handle any other exceptionsc
      print("An error occurred:", str(e))
      result = financial_crew.run()

  
  elapsed_time = time.time() - st
  elapsed_time_minutes = time.strftime("%H:%M:%S", time.gmtime(elapsed_time))
  print('Execution time:', elapsed_time_minutes)

  report = financial_crew.introduction.output.result() + '\n'+ financial_crew.impact_human.output.result() +'\n'+ financial_crew.governance.output.result() +'\n' + financial_crew.mandiant.output.result() +'\n'+financial_crew.conclusion.output.result() + '\n' + "References" + "\n 1." + str(financial_crew.ref_list[1]) + "\n 2. " + str(financial_crew.ref_list[2]) + "\n 3. " + str(financial_crew.ref_list[3])+'\n' + "Time taken (minutes) :" + str(elapsed_time_minutes)
  # print("\n\n########################")
  # print("## Here is the Report")
  # print("########################\n")
#   print(f"""
#       START
      
#       {financial_crew.impact_human.output.result()}
     
#       END
# """)
  #  {financial_crew.governance.output.result()}
  #     {financial_crew.mandiant.output.result()}
  #{financial_crew.introduction.output.result},
  #{financial_crew.conclusion.output.result}

  from docx import Document
  from docx.shared import Pt
  from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
  from docx.oxml.ns import nsdecls
  from docx.oxml import parse_xml

  def add_custom_heading(document, text, level):
      heading = document.add_heading(level=level)
      heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
      run = heading.add_run(text)
      run.bold = True

  def add_custom_paragraph(document, text, bold=False, alignment=None):
      paragraph = document.add_paragraph()
      if alignment:
          paragraph.alignment = alignment
      run = paragraph.add_run(text)
      if bold:
          run.bold = True

  def process_paragraphs(paragraphs):
      processed_paragraphs = []
      #to process the output from gemini pro llm
      #the different characters at the start of sentences of the output from Gemini Pro
      #allows us to differentiate the sentences from headers
      for line in paragraphs.split('\n'):
          if (line.strip().startswith('**') and line.strip().endswith(':**')):
              processed_paragraphs.append(('header', 3, line.strip().replace("*","")))
          elif line.strip().startswith('##'):
              processed_paragraphs.append(('header', 2, line.strip().replace("#","")))
          elif (line.strip().startswith('**') or line.strip().endswith('**')):
              processed_paragraphs.append(('header', 4, line.strip().replace("*","")))
          elif line.strip().startswith('-') or line.strip().startswith('*') :
              processed_paragraphs.append(('normal', 0, line.strip().replace("*","")))
          elif not(line.strip()):
              pass
          else:
              #other normal lines without any special characters at the start
              #to account for outputs that are from other LLMs and not from Gemini Pro
              processed_paragraphs.append(('normal', 0, line.strip()))
      return processed_paragraphs

  doc = Document()

  # Apply default style to the document
  default_style = doc.styles['Normal']
  default_font = default_style.font
  default_font.bold = False  # Change to True if you want default text to be bold
  default_font.size = Pt(12)  # Change the font size as needed

  processed_paragraphs = process_paragraphs(report)

  for paragraph_type, level, text in processed_paragraphs:
      if paragraph_type == 'header':
          add_custom_heading(doc, text, level)
      elif paragraph_type == 'normal':
          add_custom_paragraph(doc, text, bold=False)
      # else:
      #     if '**' in text:
      #         text = text.replace('**', '')
      #         add_custom_paragraph(doc, text, bold=True)
      else:
          pass

  import datetime

  current_time = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

  doc.save("AI_Report_" + current_time + ".docx")

    


    
